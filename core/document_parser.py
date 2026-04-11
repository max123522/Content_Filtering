"""
document_parser.py — Extract plain text from PDF and Word documents.

Supported formats:
    * .pdf   — via PyMuPDF (fitz)
    * .docx  — via python-docx
    * .doc   — conversion attempt via python-docx (limited)

Text is split into overlapping segments for context-aware analysis.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Union

import fitz  # PyMuPDF
import docx

import config


class ExtractionError(Exception):
    """Raised when a document cannot be parsed."""


def extract_text(source: Union[str, Path, bytes], filename: str = "") -> str:
    """
    Extract all text from a document file or bytes buffer.

    Args:
        source:   File path (str or Path) or raw bytes of the document.
        filename: Original filename (used to determine format when source is bytes).

    Returns:
        Full UTF-8 text content.

    Raises:
        ExtractionError: If the file format is unsupported or parsing fails.
    """
    if isinstance(source, (str, Path)):
        path = Path(source)
        ext = path.suffix.lower()
        with open(path, "rb") as fh:
            raw_bytes = fh.read()
    else:
        raw_bytes = source
        ext = Path(filename).suffix.lower() if filename else ""

    if ext in (".pdf",):
        return _extract_pdf(raw_bytes)
    elif ext in (".docx", ".doc"):
        return _extract_docx(raw_bytes)
    else:
        raise ExtractionError(
            f"Unsupported file format '{ext}'. "
            "Supported formats: .pdf, .docx, .doc"
        )


def _extract_pdf(raw: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF."""
    try:
        doc = fitz.open(stream=raw, filetype="pdf")
        pages: list[str] = []
        for page in doc:
            pages.append(page.get_text("text"))
        doc.close()
        return "\n".join(pages)
    except Exception as exc:
        raise ExtractionError(f"PDF extraction failed: {exc}") from exc


def _extract_docx(raw: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        doc = docx.Document(io.BytesIO(raw))
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        # Also extract table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)
        return "\n".join(paragraphs)
    except Exception as exc:
        raise ExtractionError(f"DOCX extraction failed: {exc}") from exc


def segment_text(
    text: str,
    window_size: int = config.SEGMENT_WINDOW_SIZE,
    overlap: int = 50,
) -> list[str]:
    """
    Split *text* into overlapping word-level windows.

    Args:
        text:        Full document text.
        window_size: Number of words per segment.
        overlap:     Words shared between consecutive segments.

    Returns:
        List of text segments.
    """
    words = text.split()
    if not words:
        return []

    segments: list[str] = []
    step = max(1, window_size - overlap)
    for start in range(0, len(words), step):
        chunk = words[start : start + window_size]
        segments.append(" ".join(chunk))
        if start + window_size >= len(words):
            break
    return segments


def extract_and_segment(
    source: Union[str, Path, bytes],
    filename: str = "",
) -> tuple[str, list[str]]:
    """
    Extract text and return both the full text and its segmented form.

    Returns:
        (full_text, segments)
    """
    full_text = extract_text(source, filename)
    segments = segment_text(full_text)
    return full_text, segments
